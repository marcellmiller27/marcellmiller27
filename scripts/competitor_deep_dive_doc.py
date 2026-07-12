# JHI-SIG: 69M2705M | Competitor Deep-Dive (Word) | John Henry Investments (proprietary)
"""Generate the professional Word (.docx) competitor deep-dive / reverse-engineering audit.

Targets: Mergr, S&P Global Market Intelligence (Capital IQ Pro), and CB Insights.
Purpose: find the solvable "diamond in the rough" pain point and give the Board a
cost / risk / reward read with a recommendation.

All competitor figures are best-available PUBLIC estimates and must be verified
(via live trial + written quotes) before any board decision — flagged throughout.

Run:  python scripts/competitor_deep_dive_doc.py [output.docx]
"""

from __future__ import annotations

import sys
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

SIG = "JHI-SIG: 69M2705M"
ENTITY = "JHI Research & Analytics Firm, Inc."
NAVY = RGBColor(0x0C, 0x1F, 0x33)
GOLD = RGBColor(0x9A, 0x6B, 0x12)
MUTED = RGBColor(0x5A, 0x6B, 0x7D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _shade(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _cell_text(cell, text: str, *, bold: bool = False, color: RGBColor | None = None, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else NAVY


def _muted(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def _bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        # allow a leading "**Bold:** rest" convention
        if it.startswith("**") and "**" in it[2:]:
            head, rest = it[2:].split("**", 1)
            r = p.add_run(head)
            r.bold = True
            p.add_run(rest)
        else:
            p.add_run(it)


def _table(doc: Document, headers: list[str], rows: list[list[str]], *, col_widths: list[float] | None = None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        _cell_text(hdr[i], h, bold=True, color=WHITE, size=9)
        _shade(hdr[i], "0C1F33")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            _cell_text(cells[i], val, bold=(i == 0 and len(headers) > 2), size=9)
    return t


def _footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"© {date.today().year} {ENTITY}  ·  {SIG}  ·  Confidential — not for redistribution")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def build() -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ---- Title block ----
    eyebrow = doc.add_paragraph()
    r = eyebrow.add_run(ENTITY.upper())
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = GOLD

    title = doc.add_paragraph()
    tr = title.add_run("Competitive Deep-Dive & Reverse-Engineering Audit")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sr = sub.add_run("Mergr · S&P Global Market Intelligence (Capital IQ Pro) · CB Insights")
    sr.font.size = Pt(13)
    sr.font.color.rgb = NAVY

    meta = doc.add_paragraph()
    mr = meta.add_run(
        f"Confidential — Board Discussion Draft  ·  {date.today().isoformat()}  ·  "
        f"Prepared by Cy (VP, Software Engineering — AI teammate)  ·  {SIG}"
    )
    mr.italic = True
    mr.font.size = Pt(9)
    mr.font.color.rgb = MUTED

    _muted(
        doc,
        "Not legal, tax, or investment advice. Competitor figures are best-available PUBLIC "
        "estimates (product pages, pricing intelligence, user reviews) and must be verified via "
        "live trial and written quotes before any board decision. This document is intended for "
        "the JHI Board of Directors and is a companion to docs/COMPETITOR_TEARDOWN_AND_GAP_MAP.md "
        "and docs/COMPETITOR_DEEP_DIVE_PAIN_POINTS.md.",
    )

    # ---- 1. Executive summary ----
    _heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "We reverse-engineered three market-intelligence platforms — Mergr, S&P Global Market "
        "Intelligence (Capital IQ Pro), and CB Insights — to find a solvable, underserved pain "
        "point (\u201cdiamond in the rough\u201d) that JHI can win at acceptable cost and risk. The "
        "goal is not to out-data the incumbents; it is to identify where their revenue mechanics "
        "leave a defensible gap that our end-to-end diligence workflow already addresses."
    )
    _bullets(
        doc,
        [
            "**Finding:** All three sell expensive data or signals; none deliver an affordable, "
            "transparent, decision-ready diligence workflow for the SMB / lower-middle-market "
            "acquirer (searcher, independent sponsor, family office, RIA, business broker).",
            "**The diamond:** the automated first-pass CIM read \u2192 normalized earnings \u2192 "
            "CPA-signed Quality of Earnings \u2192 close, priced for the buyer the incumbents design "
            "out. This is JHI\u2019s Deal X-Ray + QoE + Pipeline, exactly.",
            "**Verdict:** WATCH \u2192 GO on a narrow beachhead, conditional on (a) validating the "
            "Deal/Opportunity Score, (b) securing NASDAQ/Sharadar derived-data distribution rights, "
            "and (c) holding costs lean. Do NOT attempt to match S&P or CB Insights on data breadth.",
        ],
    )

    # ---- 2. Method ----
    _heading(doc, "2. Reverse-Engineering Method", 1)
    doc.add_paragraph(
        "For each platform we extract seven lenses, consistent with our existing teardown rubric: "
        "positioning/ICP, information architecture, the \u201cplumbing\u201d (data + models + tools "
        "+ outputs), moat, validation/trust strategy, pricing & packaging, and the open pain points "
        "we can exploit. We separate revenue mechanics from aesthetics and compete only where we can "
        "win: integration, price, the underserved niche, and radical transparency."
    )

    # ---- 3. Per-competitor teardowns ----
    _heading(doc, "3. Competitor Teardowns", 1)

    _heading(doc, "3.1 Mergr — the accessible M&A deal directory", 2)
    _bullets(
        doc,
        [
            "**Snapshot:** Independent M&A / private-equity transaction database; company, acquirer, "
            "and PE-firm profiles with add-on histories. Clean, honest \u201cWho uses it\u201d framing.",
            "**ICP:** corporate development, M&A advisors, business brokers, PE business-development, "
            "and searchers doing market mapping / comparable-deal research.",
            "**Plumbing (data \u2192 output):** aggregated from public sources (press releases, filings, "
            "news). Broad on announced deals; shallow on underlying financials. No proprietary score; "
            "output is directory profiles, deal lists, and ~15 pre-built M&A workflows (Buyer Match, "
            "Dossier, Roll-up tracker).",
            "**Moat:** modest \u2014 tidy UX, transparent low price, and clear positioning. Not a deep "
            "data or workflow moat; the most emulatable of the three.",
            "**Pricing (verified, public):** Mergr Pro $150/mo (~$1,800/yr) + $100/user/mo for team "
            "seats; 7-day free self-serve trial, month-to-month, cancel anytime; custom/annual plans "
            "from 10 seats. Positions itself explicitly as a PitchBook alternative "
            "($150/mo vs. $25k\u2013$50k+/yr).",
            "**Pain points / gaps:** not a diligence tool; thin financial depth; no CIM ingestion, "
            "no normalized earnings, no QoE, no close workflow.",
            "**JHI angle:** copy Mergr\u2019s transparency and \u201cWho it\u2019s for\u201d clarity; go "
            "beyond it by owning the diligence workflow it never attempts.",
        ],
    )

    _heading(doc, "3.2 S&P Global Market Intelligence (Capital IQ Pro) — the financial-breadth king", 2)
    _bullets(
        doc,
        [
            "**Snapshot:** Institutional financial-data platform (successor to Capital IQ). Deep "
            "public + private financials, estimates, transactions, ownership, and S&P credit data; "
            "Excel plug-in, screening, and Kensho/ChatIQ AI layers.",
            "**ICP:** investment banking, private equity, corporate development, equity & credit "
            "research, and large corporates.",
            "**Plumbing:** tens of millions of company records, standardized financials, an "
            "industry-standard Excel add-in, and integrated S&P ratings/credit \u2014 a 10M+ "
            "research-hour-class data operation.",
            "**Moat:** very high \u2014 data breadth + S&P brand/ratings + entrenched Excel workflow "
            "lock-in + enterprise integrations. Not displaceable on data.",
            "**Pricing (no public list; third-party estimates):** per-user tiers roughly Essentials "
            "$12k / Standard $20k / Advanced $25k per year; Vendr procurement (55 buys) shows org "
            "contracts ~$15k\u2013$215k/yr, median ~$53k; total cost of ownership often ~70% above list "
            "(integrations, support). Annual, auto-renewing.",
            "**Pain points / gaps:** cost; steep learning curve; overkill for SMB buyers; per-seat "
            "gouging and opaque pricing; no Main-Street SMB deal-diligence layer.",
            "**JHI angle:** do not compete on breadth. Win on price, approachability, and the "
            "end-to-end SMB diligence workflow Capital IQ is not built to serve.",
        ],
    )

    _heading(doc, "3.3 CB Insights — the private-market signals & content engine", 2)
    _bullets(
        doc,
        [
            "**Snapshot:** Tech / venture / private-market intelligence \u2014 startup & funding data, "
            "patents, news, market maps, the proprietary \u201cMosaic\u201d company-health score, and a "
            "large analyst-content/newsletter brand.",
            "**ICP:** corporate strategy & innovation teams, venture capital, and tech-focused M&A / "
            "business development.",
            "**Plumbing:** proprietary signals (Mosaic score), market maps, analyst reports, and an "
            "AI query layer; a strong content flywheel (newsletter reach) drives top-of-funnel.",
            "**Moat:** proprietary scoring + analyst content brand + newsletter distribution + network "
            "effects. Content flywheel is the standout, emulatable mechanic.",
            "**Pricing (no public list; third-party estimates):** annual, custom-quoted; Vendr median "
            "~$47k/yr, typical ~$60k/yr, enterprise/Insider $120k\u2013$265k+/yr, floor ~$50k. No free "
            "self-serve tier or full-platform trial (free = research content + newsletter + a 10-day "
            "trial on request).",
            "**Pain points / gaps:** expensive; tech/startup-centric (not SMB Main-Street "
            "acquisitions); cannot ingest your own CIM/documents; search-UX complaints; no QoE or "
            "close workflow; ROI questioned by some buyers.",
            "**JHI angle:** emulate the transparent-methodology + content-flywheel mechanics; aim the "
            "score at SMB deal decisions CB Insights does not serve, and let users bring their own CIM.",
        ],
    )

    # ---- 4. Cross-competitor synthesis ----
    _heading(doc, "4. Cross-Competitor Synthesis", 1)
    doc.add_paragraph("Scored 1 (weak) \u2013 5 (strong) for the SMB/search-fund acquirer\u2019s lens:")
    _table(
        doc,
        ["Platform", "SMB fit", "Data depth", "Diligence workflow", "Price access", "Primary moat", "Biggest gap"],
        [
            ["Mergr", "3", "2", "1", "5", "Price + clean UX", "No diligence layer"],
            ["S&P Capital IQ", "2", "5", "1", "1", "Data breadth + brand", "Overkill / costly for SMB"],
            ["CB Insights", "2", "4", "1", "1", "Signals + content brand", "Tech-centric, no CIM/QoE"],
            ["JHI (target)", "5", "3", "5", "5", "End-to-end + transparency", "Score validation pending"],
        ],
    )
    _muted(
        doc,
        "Read: every incumbent scores 1 on diligence workflow \u2014 the column JHI is built to own. "
        "Data depth is their moat, not ours; we win on the workflow, price, and niche they ignore.",
    )

    # ---- 5. Diamond in the rough ----
    _heading(doc, "5. The \u201cDiamond in the Rough\u201d \u2014 Opportunity Thesis", 1)
    doc.add_paragraph(
        "The recurring, solvable pain across all three is the same one their own users describe: for "
        "the SMB / lower-middle-market acquirer, the bottleneck is not judgment \u2014 it is reading. "
        "A first-pass review of a data room, tying the CIM to the financials, normalizing how the "
        "seller presented the numbers, and building a question list is days of senior time per target, "
        "most of it spent on deals that die."
    )
    _bullets(
        doc,
        [
            "**The gap:** between free/thin tools and $13k\u2013$100k+/seat institutional platforms, "
            "there is no affordable, transparent, decision-ready diligence workflow for this buyer.",
            "**The unmet job:** \u201cget to no fast\u201d \u2014 cheap early screening that kills weak "
            "deals before the $20k\u2013$30k QoE spend, then an affordable CPA-signed QoE for the "
            "survivors.",
            "**Why JHI can win it:** Deal X-Ray (first-pass CIM read + normalized earnings), a "
            "CPA-partner QoE, a published/validated score, and a close pipeline \u2014 an end-to-end "
            "workflow no incumbent offers, at a transparent price ($50 / $299 / $1,500).",
            "**Trust wedge:** publish methodology and IC/t-stats, and disclose limits \u2014 the "
            "opposite of the incumbents\u2019 black-box scores.",
        ],
    )

    # ---- 6. Cost / risk / reward ----
    _heading(doc, "6. Cost / Risk / Reward Assessment", 1)
    _table(
        doc,
        ["Dimension", "Assessment"],
        [
            ["Cost to build/serve", "Incremental. Core workflow largely built. Main run-rate: SF1 data "
             "license (~$18k/yr), cloud/AI infra, and a CPA-partner network (revenue-share, not fixed)."],
            ["Risk \u2014 data rights", "NASDAQ/Sharadar derived-data distribution rights are not yet "
             "confirmed (addendum pending, sign-by ~05-Aug-26). Gating for client-facing score serving."],
            ["Risk \u2014 validation", "The Opportunity/Deal Score is not yet validated (H5 currently "
             "FAIL on price-only factors). Serving a score we cannot defend is a trust and liability risk."],
            ["Risk \u2014 competitive", "Incumbents could move down-market; but their cost structure and "
             "enterprise DNA make a credible SMB self-serve product unlikely near-term."],
            ["Risk \u2014 positioning", "Must stay research/decision-support (software/data publisher, "
             "NAICS 513210), not investment advice or brokerage. Keep compliance framing tight."],
            ["Reward", "Underserved, growing ETA/search-fund + SMB M&A segment; pricing power via "
             "transparency; a defensible end-to-end moat; and a small data-network effect from labeled "
             "deal outcomes over time."],
            ["Risk-adjusted verdict", "WATCH \u2192 GO on a narrow beachhead, conditional on score "
             "validation + data rights + lean cost. High reward relative to modest incremental cost."],
        ],
        col_widths=[1.6, 5.0],
    )

    # ---- 7. Recommendation ----
    _heading(doc, "7. Recommendation to the Board", 1)
    _bullets(
        doc,
        [
            "**Proceed \u2014 conditionally.** Pursue the SMB/search-fund diligence wedge as the "
            "beachhead; do not chase incumbent data breadth.",
            "**Gate on validation.** Do not market the score as predictive until it clears the "
            "pre-registered bar; keep publishing results (including FAILs) as a differentiator.",
            "**Secure data rights first.** Close the NASDAQ/Sharadar derived-data addendum before "
            "any client-facing score serving.",
            "**Emulate the mechanics, not the look:** CB Insights\u2019 content flywheel, Morningstar-style "
            "published methodology, Mergr\u2019s transparency, and the incumbents\u2019 export/workflow "
            "lock-in (interactive Excel/PDF).",
            "**Keep it lean.** Hold run-rate near the current base until the score validates and the "
            "beachhead shows pull.",
        ],
    )

    _heading(doc, "Proposed board resolution (for minutes)", 2)
    p = doc.add_paragraph()
    rr = p.add_run(
        "RESOLVED, that management pursue the SMB / lower-middle-market diligence wedge identified in "
        "this audit as the Company\u2019s competitive beachhead; that no predictive claim be made for "
        "the Opportunity/Deal Score until it clears the pre-registered validation bar; that the "
        "NASDAQ/Sharadar derived-data distribution rights be secured before client-facing score "
        "serving; and that operating run-rate be held lean pending validation and demonstrated "
        "market pull. A live-trial teardown pass to verify all pricing/coverage estimates herein is "
        "authorized."
    )
    rr.italic = True

    # ---- 8. Verified pricing & sources ----
    _heading(doc, f"8. Verified Pricing & Sources (public, as of {date.today().isoformat()})", 1)
    doc.add_paragraph(
        "Pricing below is verified against current public pages and third-party procurement data "
        "(Vendr, Capchase) rather than a paid trial. Mergr publishes a list price; S&P and CB Insights "
        "do not, so their figures are sourced third-party estimates and procurement medians."
    )
    _table(
        doc,
        ["Platform", "Public list?", "Entry / typical", "Enterprise / high end", "Trial"],
        [
            ["Mergr", "Yes", "$150/mo (~$1,800/yr) + $100/seat", "Custom, 10+ seats", "7-day self-serve"],
            ["S&P Capital IQ Pro", "No", "~$12k\u2013$25k/user/yr; org median ~$53k", "up to ~$215k+/yr; TCO ~+70%", "Sales demo"],
            ["CB Insights", "No", "median ~$47k/yr; typical ~$60k", "$120k\u2013$265k+/yr", "10-day (on request); no self-serve"],
        ],
    )
    doc.add_paragraph("Sources (accessed for this pass):")
    _bullets(
        doc,
        [
            "**Mergr:** mergr.com/pricing; docs.mergr.com (\u201cPitchBook alternative\u201d; seat add-ons); "
            "Fintalent listing.",
            "**S&P Capital IQ Pro:** CostBench (Vendr procurement, 55 buys); Wall Street Prep "
            "(pricing overview); GeminIQ (per-tier estimates).",
            "**CB Insights:** EasyVC and Elevated Signal (Vendr/Capchase medians & tiers); CostBench "
            "calculator; Prospeo (CB Insights vs. Crunchbase).",
        ],
    )
    _heading(doc, "What a true paid live-trial still adds (founder action)", 2)
    _bullets(
        doc,
        [
            "**Cheap & fast:** Mergr is self-serve at $150/mo with a 7-day free trial \u2014 a real "
            "hands-on teardown is low-cost and immediate.",
            "**Gated:** S&P Capital IQ Pro and CB Insights require sales demos (no self-serve); a paid "
            "eval needs a business email, a sales conversation, and likely a signed order \u2014 founder "
            "action.",
            "**What it adds:** the pricing above is already well-sourced for board triage; a paid trial "
            "firms up first-run funnel, workflow depth, and export/plumbing detail, plus first-hand "
            "churn/pain mining \u2014 not the price, which is settled here.",
        ],
    )

    # ---- 9. Appendix ----
    _heading(doc, "9. Appendix \u2014 Sources, Ethics & Disclaimers", 1)
    _bullets(
        doc,
        [
            "**Sources:** public product/pricing pages, third-party pricing intelligence, and public "
            "user reviews (e.g., G2, TrustRadius, Reddit, search-fund forums). Competitive intelligence "
            "gathered from public sources only.",
            "**Verification status:** pricing is verified against current public pages and third-party "
            "procurement data (see \u00a78); S&P and CB Insights do not publish list prices, so their "
            "figures are sourced estimates/medians. Coverage, UX, and plumbing depth still benefit from "
            "a paid hands-on trial before any public claim.",
            "**No affiliation:** JHI is not affiliated with, endorsed by, or partnered with Mergr, "
            "S&P Global, or CB Insights. Marks belong to their respective owners.",
            f"**Provenance:** founder signature of record {SIG}. \u00a9 {date.today().year} {ENTITY}. "
            "All rights reserved. Confidential.",
        ],
    )

    _footer(doc)
    return doc


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else "JHI_Competitor_Deep_Dive.docx"
    build().save(out)
    print(f"wrote {out}")
