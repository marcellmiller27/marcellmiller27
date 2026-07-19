# JHI-SIG: 69M2705M | Company Book (Word) | JHI Research & Analytics Firm, Inc. (proprietary)
"""Generate the Word (.docx) JHI Company Book — Policy, Procedures & Processes.
Mirrors docs/COMPANY_BOOK_POLICY_PROCEDURES_PROCESSES.md.

Run:  python scripts/company_book_doc.py [output.docx]
"""

from __future__ import annotations

import sys
from datetime import date

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

SIG = "JHI-SIG: 69M2705M"
ENTITY = "JHI Research & Analytics Firm, Inc."
NAVY = RGBColor(0x0C, 0x1F, 0x33)
GOLD = RGBColor(0x9A, 0x6B, 0x12)
MUTED = RGBColor(0x5A, 0x6B, 0x7D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _shade(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _h(doc, text, size=13, color=NAVY, before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def _p(doc, text, size=10.5, color=None, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text).font.size = Pt(10.5)


def _table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        _shade(c, "0C1F33")
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)


def build() -> Document:
    doc = Document()
    t = doc.add_paragraph()
    tr = t.add_run("JHI Company Book — Policy, Procedures & Processes")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = NAVY
    s = doc.add_paragraph()
    sr = s.add_run(f"{ENTITY} (proprietary)  ·  {SIG}  ·  generated {date.today().isoformat()}")
    sr.font.size = Pt(9)
    sr.font.color.rgb = MUTED
    _p(doc, "The single source of truth for how the firm operates. Living document — every "
            "new capability is documented as Policy \u2192 Procedure \u2192 Process \u2192 Job Aid "
            "before it is considered done.", italic=True, color=MUTED)

    _h(doc, "Guiding philosophy")
    _bullet(doc, "\u201cHow we do anything is how we do everything.\u201d Standards are universal, not situational.")
    _bullet(doc, "Run live by Efficiency & Effectiveness \u2014 least wasted energy, doing the right thing well.")
    _bullet(doc, "Foundation first \u2014 structure and standards precede feature work.")

    _h(doc, "Definitions")
    _table(doc, ["Term", "Meaning"], [
        ["Policy", "The rule and the why \u2014 a standing decision that governs behavior."],
        ["Procedure", "The step-by-step \u2018how\u2019 to perform a specific task correctly."],
        ["Process", "The end-to-end flow connecting procedures/steps across the firm."],
        ["Job Aid", "A quick-reference for a single task, including troubleshooting."],
    ])

    _h(doc, "How we document (the meta-procedure) \u2014 MANDATORY")
    for b in [
        "Policy \u2014 the governing rule (if new governance is implied).",
        "Procedure \u2014 the exact steps to do it.",
        "Process \u2014 where it fits in the end-to-end flow.",
        "Job Aid \u2014 quick reference + troubleshooting (docs/TroubleShoot/Job-Aids/).",
    ]:
        _bullet(doc, b)
    _p(doc, "Nothing is \u2018done\u2019 until it is documented here (or linked from here).", bold=True)

    _h(doc, "The Book \u2014 sections & source-of-truth documents")
    _table(doc, ["Section", "Authoritative doc(s)"], [
        ["1. Governance & Engineering", "docs/ENGINEERING_POLICY.md ; AGENTS.md"],
        ["2. Product & Platform", "docs/PLATFORM_IA_BLUEPRINT.md (Mirror, Core Rule, Context, Depth principles)"],
        ["3. Pricing & Billing", "docs/PRICING_BILLING_SCHEMA.md (T1 $1,500 / T2 $299 / T3 $110 list, $1,188 prepaid)"],
        ["4. Sales", "docs/SALES_COMMISSION_MODEL.md ; scripts/sales_commission_prepaid_model.py"],
        ["5. Data, Licensing & Compliance", "IA Blueprint Part J ; MARKET_DATA_VENDOR_COMPARISON.md ; COMPANY_POSTURE_AND_COMPLIANCE.md ; SECURITY_POSTURE_AND_DATA_PROTECTION.md ; docs/legal/"],
        ["6. People & Organization", "docs/JOB_DESCRIPTIONS_AND_STAFFING_REQUIREMENTS.md ; docs/AI_AGENT_TEAM_PROFILES.md"],
        ["7. Finance & Governance Records", "docs/board/ ; docs/financial_models/ ; docs/projections/ ; docs/ip/"],
        ["8. Job Aids", "docs/TroubleShoot/Job-Aids/ (Excel Download Recovery, + future aids)"],
        ["9. Reference", "docs/GLOSSARY_AND_ACRONYMS.md ; docs/internal_valuation_package/PITCH_DECK.md"],
    ])

    _h(doc, "Maintenance")
    _bullet(doc, "Living document \u2014 update whenever a policy/procedure/process is added or changed.")
    _bullet(doc, "Word edition regenerated by scripts/company_book_doc.py after edits.")
    _bullet(doc, "Reviewed at Founder working sessions; material changes noted in board minutes.")

    foot = doc.add_paragraph()
    fr = foot.add_run(f"\u00a9 {date.today().year} {ENTITY}. All rights reserved. Confidential \u2014 "
                      f"internal use. Provenance: {SIG}.")
    fr.font.size = Pt(8)
    fr.font.color.rgb = MUTED
    return doc


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else "JHI_Company_Book_Policy_Procedures_Processes.docx"
    build().save(out)
    print(f"wrote {out}")
